from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import List, Tuple

import docx
from fastapi import HTTPException, UploadFile
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document as LangChainDocument
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from pypdf import PdfReader

from config import RAGSettings
from schemas import DocumentChunk


ALLOWED_EXTENSIONS = {".pdf", ".md", ".markdown", ".txt", ".docx"}
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_FILES_ALLOWED = 10


@dataclass
class UploadedMemoryFile:
    name: str
    content: bytes
    content_type: str


@dataclass
class RAGState:
    files: List[UploadedMemoryFile] = field(default_factory=list)
    vectorstore: Chroma | None = None
    chunks_count: int = 0


class RAGService:
    def __init__(self, settings: RAGSettings) -> None:
        self.settings = settings
        if not settings.openai_api_key:
            raise HTTPException(status_code=500, detail="OPENAI_API_KEY not configured")
        self.embedding = OpenAIEmbeddings(model=self.settings.embedding_model, api_key=self.settings.openai_api_key)
        self.llm = ChatOpenAI(
            model=self.settings.openai_model,
            temperature=self.settings.temperature,
            api_key=self.settings.openai_api_key,
        )
        self.state = RAGState()


    def _validate_file(self, upload: UploadFile) -> Tuple[bool, str]:
        extension = Path(upload.filename or "").suffix.lower()
        if extension not in ALLOWED_EXTENSIONS:
            return False, "Unsupported file type"
        return True, ""

    async def store_files(self, uploads: List[UploadFile]) -> Tuple[List[str], List[str], int]:
        accepted: List[str] = []
        rejected: List[str] = []

        if len(self.state.files) + len(uploads) > MAX_FILES_ALLOWED:
            raise HTTPException(status_code=400, detail="10 file limit reached")

        for upload in uploads:
            is_valid, reason = self._validate_file(upload)
            if not is_valid:
                rejected.append(f"{upload.filename} ({reason})")
                continue

            content = await upload.read()
            if len(content) > MAX_FILE_SIZE:
                rejected.append(f"{upload.filename} (excede 10MB)")
                continue
            self.state.files.append(
                UploadedMemoryFile(name=upload.filename, content=content, content_type=upload.content_type)
            )
            accepted.append(upload.filename or "document")

        if accepted:
            self.state.vectorstore = None
            self.state.chunks_count = 0

        return accepted, rejected, len(self.state.files)

    def _load_pdf(self, file: UploadedMemoryFile) -> List[LangChainDocument]:
        reader = PdfReader(BytesIO(file.content))
        documents: List[LangChainDocument] = []
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            documents.append(
                LangChainDocument(
                    page_content=text.strip(),
                    metadata={"source": file.name, "page": idx},
                )
            )
        return documents

    def _load_docx(self, file: UploadedMemoryFile) -> List[LangChainDocument]:
        doc = docx.Document(BytesIO(file.content))
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        return [
            LangChainDocument(
                page_content=text.strip(),
                metadata={"source": file.name, "page": 1},
            )
        ]

    def _load_text_like(self, file: UploadedMemoryFile) -> List[LangChainDocument]:
        decoded = file.content.decode("utf-8", errors="ignore")
        return [
            LangChainDocument(
                page_content=decoded.strip(),
                metadata={"source": file.name, "page": 1},
            )
        ]

    def _to_documents(self, file: UploadedMemoryFile) -> List[LangChainDocument]:
        extension = Path(file.name).suffix.lower()
        if extension == ".pdf":
            return self._load_pdf(file)
        if extension == ".docx":
            return self._load_docx(file)
        if extension in {".md", ".markdown", ".txt"}:
            return self._load_text_like(file)
        raise HTTPException(status_code=400, detail=f"Unsupported extension: {extension}")

    def process_documents(self) -> int:
        if not self.state.files:
            raise HTTPException(status_code=400, detail="No files to process")

        documents: List[LangChainDocument] = []
        for file in self.state.files:
            documents.extend(self._to_documents(file))

        splitter = RecursiveCharacterTextSplitter(
            chunk_size=self.settings.chunk_size, chunk_overlap=self.settings.chunk_overlap
        )
        chunks = splitter.split_documents(documents)
        vectorstore = Chroma.from_documents(
            documents=chunks,
            embedding=self.embedding,
            collection_name="uploaded_documents",
        )
        self.state.vectorstore = vectorstore
        self.state.chunks_count = len(chunks)
        return len(chunks)

    def query(self, question: str) -> tuple[str, List[DocumentChunk]]:
        if not question.strip():
            raise HTTPException(status_code=400, detail="Question cannot be empty")
        if not self.state.vectorstore:
            raise HTTPException(status_code=400, detail="Process documents first")

        retriever = self.state.vectorstore.as_retriever(search_kwargs={"k": self.settings.k_results})
        relevant_docs: List[LangChainDocument] = retriever.invoke(question)

        context_blocks = []
        chunk_responses: List[DocumentChunk] = []
        for doc in relevant_docs:
            source = Path(doc.metadata.get("source", "")).name
            page = doc.metadata.get("page", "?")
            context_blocks.append(f"Source: {source} (page {page})\n{doc.page_content}")
            chunk_responses.append(DocumentChunk(text=doc.page_content, source=source, page=page))

        context_text = "\n\n".join(context_blocks)
        prompt = ChatPromptTemplate.from_messages(
            [
                (
                    "system",
                    "You are an expert lab operations assistant. Answer strictly from the provided context. "
                    "If the information is missing, say it is not available. "
                    "Always respond in the same language as the question. "
                    "Be concise and clear in your answers.",
                ),
                (
                    "human",
                    "Context:\n{context}\n\nQuestion: {question}\nAnswer:",
                ),
            ]
        )

        chain = prompt | self.llm
        response = chain.invoke({"context": context_text, "question": question})
        answer_text = response.content if hasattr(response, "content") else str(response)
        answer = f"{answer_text}"
        return answer, chunk_responses

    def reset(self) -> tuple[int, int]:
        files_cleared = len(self.state.files)
        chunks_cleared = self.state.chunks_count

        if self.state.vectorstore and hasattr(self.state.vectorstore, "delete_collection"):
            try:
                self.state.vectorstore.delete_collection()
            except Exception:
                # Vector store cleanup is best effort; continue to reset in-memory state.
                pass

        self.state = RAGState()
        return files_cleared, chunks_cleared
